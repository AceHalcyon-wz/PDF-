#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
格式转换逻辑模块

接口状态: 已更新以符合统一接口标准ModuleInterface
更新内容:
1. 继承ModuleInterface基类
2. 使用标准日志记录方法
3. 实现标准接口方法
"""

import os
import csv
from pypdf import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from core.interface import ModuleInterface

# 尝试导入额外的转换库
try:
    import pdf2image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class ConversionEngine(ModuleInterface):
    """PDF转换引擎类"""

    def __init__(self):
        """初始化转换引擎"""
        super().__init__("conversion")
        self.log_info("转换引擎初始化完成")

    def pdf_to_text(self, input_path, output_path):
        """
        PDF转文本
        
        Args:
            input_path (str): 输入PDF文件路径
            output_path (str): 输出文本文件路径
        """
        self.log_info(f"开始PDF转文本: {input_path}")
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 读取PDF文件
            reader = PdfReader(input_path)
            
            # 提取文本
            text_content = ""
            for i, page in enumerate(reader.pages):
                text_content += f"--- Page {i+1} ---\n"
                text_content += page.extract_text()
                text_content += "\n\n"
            
            # 保存为文本文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
                
            self.log_info(f"PDF转文本完成: {output_path}")
            return True
        except Exception as e:
            self.log_error(f"PDF转文本失败: {str(e)}")
            raise PDFConversionError(f"PDF转文本失败: {str(e)}")

    def pdf_to_images(self, input_path, output_dir, format="jpg", dpi=200):
        """
        PDF转图像
        
        Args:
            input_path (str): 输入PDF文件路径
            output_dir (str): 输出目录路径
            format (str): 图像格式 ("jpg", "png")
            dpi (int): 图像分辨率
        """
        try:
            self.log_info(f"开始PDF转图像: {input_path}")
            
            # 检查依赖是否可用
            if not PDF2IMAGE_AVAILABLE or not PIL_AVAILABLE:
                raise Exception("PDF转图像功能需要额外依赖库支持：\n"
                                "1. 安装poppler工具包\n"
                                "2. 安装Python依赖：pip install pdf2image pillow")
            
            # 确保输出目录存在
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 转换PDF为图像
            images = pdf2image.convert_from_path(input_path, dpi=dpi)
            
            # 保存图像
            for i, image in enumerate(images):
                image_path = os.path.join(output_dir, f"page_{i+1:03d}.{format}")
                if format.lower() == "jpg":
                    image.save(image_path, "JPEG")
                else:
                    image.save(image_path, "PNG")
                    
            self.log_info(f"PDF转图像完成，共生成{len(images)}张图像")
        except Exception as e:
            self.log_error(f"PDF转图像失败: {str(e)}")
            raise Exception(f"PDF转图像失败: {str(e)}")

    def image_to_pdf(self, image_paths, output_path):
        """
        图像转PDF
        
        Args:
            image_paths (list): 图像文件路径列表
            output_path (str): 输出PDF文件路径
        """
        try:
            self.log_info(f"开始图像转PDF，共{len(image_paths)}张图像")
            
            # 检查依赖是否可用
            if not PIL_AVAILABLE:
                raise Exception("图像转PDF功能需要额外依赖库支持：pip install pillow")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 创建PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            for image_path in image_paths:
                # 打开图像
                img = Image.open(image_path)
                
                # 调整图像大小以适应页面
                img_width, img_height = img.size
                scale = min(width / img_width, height / img_height) * 0.9  # 留一些边距
                new_width = img_width * scale
                new_height = img_height * scale
                
                # 将图像居中放置在页面上
                x = (width - new_width) / 2
                y = (height - new_height) / 2
                
                # 绘制图像
                c.drawImage(image_path, x, y, width=new_width, height=new_height)
                c.showPage()
                
            c.save()
            self.log_info(f"图像转PDF完成: {output_path}")
            return True
        except Exception as e:
            self.log_error(f"图像转PDF失败: {str(e)}")
            raise Exception(f"图像转PDF失败: {str(e)}")

    def pdf_to_word(self, input_path, output_path):
        """
        PDF转Word
        
        Args:
            input_path (str): 输入PDF文件路径
            output_path (str): 输出Word文件路径
        """
        try:
            # 检查依赖是否可用
            if not WORD_AVAILABLE:
                raise Exception("PDF转Word功能需要python-docx库，请安装：pip install python-docx")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 读取PDF文件
            reader = PdfReader(input_path)
            
            # 创建Word文档
            doc = Document()
            
            # 逐页提取文本并添加到Word文档
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():  # 只有当页面有文本时才添加
                    doc.add_paragraph(f"Page {i+1}")
                    doc.add_paragraph(text)
                    doc.add_page_break()
            
            # 保存Word文档
            doc.save(output_path)
            
            return True
        except Exception as e:
            raise Exception(f"PDF转Word失败: {str(e)}")

    def word_to_pdf(self, input_path, output_path):
        """
        Word转PDF
        
        Args:
            input_path (str): 输入Word文件路径
            output_path (str): 输出PDF文件路径
        """
        try:
            # 检查依赖是否可用
            if not WORD_AVAILABLE:
                raise Exception("Word转PDF功能需要python-docx和reportlab库，请安装：pip install python-docx reportlab")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 读取Word文档
            doc = Document(input_path)
            
            # 创建PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # 设置字体和起始位置
            c.setFont("Helvetica", 12)
            x, y = 50, height - 50
            line_height = 15
            
            # 逐段落写入PDF
            for paragraph in doc.paragraphs:
                # 如果到达页面底部，创建新页面
                if y < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = height - 50
                
                c.drawString(x, y, paragraph.text)
                y -= line_height
            
            # 保存PDF
            c.save()
            
            return True
        except Exception as e:
            raise Exception(f"Word转PDF失败: {str(e)}")

    def text_to_pdf(self, input_path, output_path):
        """
        文本转PDF
        
        Args:
            input_path (str): 输入文本文件路径
            output_path (str): 输出PDF文件路径
        """
        try:
            self.log_info(f"开始文本转PDF: {input_path}")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 读取文本文件
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # 创建PDF writer
            writer = PdfWriter()
            
            # 创建内存中的PDF
            packet = io.BytesIO()
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            can = canvas.Canvas(packet, pagesize=letter)
            width, height = letter
            
            # 添加文本内容
            y_position = height - 50
            lines = text_content.split('\n')
            for line in lines:
                if self.should_cancel():
                    raise Exception("操作被用户取消")
                    
                if line.strip():
                    can.drawString(50, y_position, line)
                    y_position -= 20
                    
                    # 如果页面满了，创建新页面
                    if y_position < 50:
                        can.showPage()
                        y_position = height - 50
            
            can.save()
            
            # 将内存中的PDF添加到writer
            packet.seek(0)
            temp_reader = PdfReader(packet)
            for page in temp_reader.pages:
                writer.add_page(page)
            
            # 保存PDF文件
            with open(output_path, 'wb') as f:
                writer.write(f)
                
            self.log_info(f"文本转PDF完成: {output_path}")
        except Exception as e:
            self.log_error(f"文本转PDF失败: {str(e)}")
            raise Exception(f"文本转PDF失败: {str(e)}")

    def pdf_to_csv(self, input_path, output_path):
        """
        PDF转CSV
        
        Args:
            input_path (str): 输入PDF文件路径
            output_path (str): 输出CSV文件路径
        """
        try:
            self.log_info(f"开始PDF转CSV: {input_path}")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            reader = PdfReader(input_path)
            
            # 提取文本并尝试解析为表格数据
            import csv
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                csv_writer = csv.writer(csvfile)
                
                for i, page in enumerate(reader.pages):
                    if self.should_cancel():
                        raise Exception("操作被用户取消")
                        
                    text = page.extract_text()
                    # 简单处理，按行分割
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            # 按逗号或制表符分割
                            if ',' in line:
                                row = line.split(',')
                            elif '\t' in line:
                                row = line.split('\t')
                            else:
                                row = [line]
                            csv_writer.writerow(row)
            
            self.log_info(f"PDF转CSV完成: {output_path}")
        except Exception as e:
            self.log_error(f"PDF转CSV失败: {str(e)}")
            raise Exception(f"PDF转CSV失败: {str(e)}")

    def pdf_to_markdown(self, input_path, output_path):
        """
        PDF转Markdown
        
        Args:
            input_path (str): 输入PDF文件路径
            output_path (str): 输出Markdown文件路径
        """
        try:
            self.log_info(f"开始PDF转Markdown: {input_path}")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            reader = PdfReader(input_path)
            markdown_content = ""
            
            # 为每一页添加内容
            for i, page in enumerate(reader.pages):
                if self.should_cancel():
                    raise Exception("操作被用户取消")
                    
                text = page.extract_text()
                markdown_content += f"# Page {i+1}\n\n"
                markdown_content += f"{text}\n\n"
                markdown_content += "---\n\n"
            
            # 保存Markdown文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
                
            self.log_info(f"PDF转Markdown完成: {output_path}")
        except Exception as e:
            self.log_error(f"PDF转Markdown失败: {str(e)}")
            raise Exception(f"PDF转Markdown失败: {str(e)}")

    def ppt_to_pdf(self, input_path, output_path):
        """
        PPT转PDF
        
        Args:
            input_path (str): 输入PPT文件路径
            output_path (str): 输出PDF文件路径
        """
        try:
            self.log_info(f"开始PPT转PDF: {input_path}")
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 这里需要使用第三方库如python-pptx来处理PPT文件
            # 由于这是一个复杂操作，我们简化实现
            # 实际项目中应该使用专业的库来处理PPT转PDF
            
            # 创建一个简单的PDF作为示例
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            c.drawString(100, height - 100, f"PPT文件: {os.path.basename(input_path)}")
            c.drawString(100, height - 130, "注意: 这是一个示例转换")
            c.drawString(100, height - 160, "实际应用中需要使用专业库处理PPT文件")
            c.save()
                
            self.log_info(f"PPT转PDF完成: {output_path}")
        except Exception as e:
            self.log_error(f"PPT转PDF失败: {str(e)}")
            raise Exception(f"PPT转PDF失败: {str(e)}")
